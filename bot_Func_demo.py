from aiogram import Dispatcher, types, F, Bot
from aiogram.types import Message
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from aiogram.filters.command import Command
from config import TOKEN
import re
import asyncio
import datetime


def find_first_match(pattern, text, default="Нет данных"):
    matches = re.findall(pattern, text)
    return matches[0] if matches else default


class ReportData:
    def __init__(self):
        self.time_departure = None
        self.odometer_reading = None
        self.take_picture_new_odometer = None
        self.arrival_object = None
        self.arrival_complex = None
        self.arrival_time = None
        self.actions_taken = None
        self.departure_time = None
        self.intermediate_odometer = None
        self.take_picture_intermediate_odometer = None
        self.final_odometer = None
        self.total_distance = None
        self.time_spent = None

    def fill_report(self, data):
        # Обновляем данные только если они присутствуют в переданном словаре
        if 'time_departure' in data:
            self.time_departure = data['time_departure']
        if 'odometer_reading' in data:
            self.odometer_reading = data['odometer_reading']
        if 'take_picture_new_odometer' in data:
            self.take_picture_new_odometer = data['take_picture_new_odometer']
        if 'arrival_object' in data:
            self.arrival_object = data['arrival_object']
        if 'arrival_complex' in data:
            self.arrival_complex = data['arrival_complex']
        if 'arrival_time' in data:
            self.arrival_time = data['arrival_time']
        if 'actions_taken' in data:
            self.actions_taken = data['actions_taken']
        if 'departure_time' in data:
            self.departure_time = data['departure_time']
        if 'intermediate_odometer' in data:
            self.intermediate_odometer = data['intermediate_odometer']
        if 'take_picture_intermediate_odometer' in data:
            self.take_picture_intermediate_odometer = data['take_picture_intermediate_odometer']
        if 'final_odometer' in data:
            self.final_odometer = data['final_odometer']

    def get_report(self):
        # Возвращает отчет в виде словаря
        return {
            'time_departure': self.time_departure,
            'odometer_reading': self.odometer_reading,
            'take_picture_new_odometer': self.take_picture_new_odometer,
            'arrival_object': self.arrival_object,
            'arrival_complex': self.arrival_complex,
            'arrival_time': self.arrival_time,
            'actions_taken': self.actions_taken,
            'departure_time': self.departure_time,
            'intermediate_odometer': self.intermediate_odometer,
            'take_picture_intermediate_odometer': self.take_picture_intermediate_odometer,
            'final_odometer': self.final_odometer
            # Добавьте остальные переменные, если есть
        }


def total_distance(start_reading, end_reading):
    print("Start reading:", start_reading)
    print("End reading:", end_reading)
    try:
        return int(end_reading) - int(start_reading)
    except (TypeError, ValueError):
        return None


def time_odds(start_time, end_time):
    print("Start time:", start_time)
    print("End time:", end_time)
    try:
        start_hour, start_minute = map(int, start_time.split())
        end_hour, end_minute = map(int, end_time.split())
        total_hours = end_hour - start_hour
        total_minutes = end_minute - start_minute
        if total_minutes < 0:
            total_hours -= 1
            total_minutes += 60
        return f"{total_hours} часов {total_minutes} минут"
    except (AttributeError, ValueError, TypeError):
        return None


def create_report(report_data, report_date):
    # Создание нового документа
    document = Document()

    # Настройки документа
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal'].font.size = Pt(14)

    # Заголовок
    title = document.add_paragraph()
    title_run1 = title.add_run('Отчет по АВР')
    title_run1.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run1.style.font.size = Pt(16)
    title.paragraph_format.space_after = Pt(0)  # Установка интервала "После" в 0pt

    title2 = document.add_paragraph()
    title_run2 = title2.add_run(f'за {report_date}.')
    title_run2.bold = True
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2.style.font.size = Pt(16)

    # Добавление пустой строки
    document.add_paragraph()

    # Добавление данных из класса ReportData
    start_time = document.add_paragraph()
    start_time.add_run('Время выезда: ').bold = True
    start_time.add_run(report_data.data.get('Время выезда', ''))
    start_time.paragraph_format.space_after = Pt(0)

    start_odometer = document.add_paragraph()
    start_odometer.add_run('Показания одометра: ').bold = True
    start_odometer.add_run(report_data.data.get('Показания одометра', ''))
    start_odometer.paragraph_format.space_after = Pt(0)

    pribytie_text = document.add_paragraph()
    pribytie_text.add_run('Прибытие на заявку: ').bold = True
    pribytie_text.add_run(f"{report_data.data.get('Прибытие на заявку', '')} комплекс {report_data.data.get('Комплекс', '')}")
    pribytie_text.paragraph_format.space_after = Pt(0)

    object_time = document.add_paragraph()
    object_time.add_run('Время прибытия: ').bold = True
    object_time.add_run(report_data.data.get('Время прибытия', ''))
    object_time.paragraph_format.space_after = Pt(0)

    action = document.add_paragraph()
    action.add_run('Предпринятые действия: ').bold = True
    action.add_run(', '.join(report_data.data.get('Предпринятые действия', '')))
    action.paragraph_format.space_after = Pt(0)

    final_time = document.add_paragraph()
    final_time.add_run('Время убытия: ').bold = True
    final_time.add_run(report_data.data.get('Время убытия', ''))
    final_time.paragraph_format.space_after = Pt(0)

    cost_time = document.add_paragraph()
    cost_time.add_run('Время пребывания на месте работы: ').bold = True
    cost_time.add_run(f"{report_data.data.get('Время пребывания на месте работы', '')} часа")
    cost_time.paragraph_format.space_after = Pt(0)

    inter_odometer = document.add_paragraph()
    inter_odometer.add_run('Промежуточные показания одометра: ').bold = True
    inter_odometer.add_run(report_data.data.get('Промежуточные показания одометра', ''))
    inter_odometer.paragraph_format.space_after = Pt(0)

    final_odometer = document.add_paragraph()
    final_odometer.add_run('Показания одометра: ').bold = True
    final_odometer.add_run(report_data.data.get('Показания одометра', ''))
    final_odometer.paragraph_format.space_after = Pt(0)

    distance = document.add_paragraph()
    distance.add_run('Пройдено расстояния всего: ').bold = True
    distance.add_run(f"{report_data.data.get('Пройдено расстояния всего', '')} км.")
    distance.paragraph_format.space_after = Pt(0)

    # Сохранение документа
    document.save(f'Отчет по АВР на {report_date}.docx')


# Создание объектов бота и диспетчера
bot = Bot(token=TOKEN)
dp = Dispatcher(bot=bot)


# Обработчик команды /start
@dp.message(Command("start"))
async def send_welcome(message: types.Message):
    await message.reply("Привет! Отправьте мне текст для создания отчета по АВР.")


# Обработчик текстового сообщения с данными для отчета
@dp.message(Command("Text"))
async def process_report_text(message: types.Message):
    text = message.text
    report_data = ReportData(text)
    report_date = datetime.datetime.now().strftime("%d %m %Y")
    await message.reply(f"Отчет по АВР за {report_date}:\n"
                        f"Время выезда: {report_data.data.get('Время выезда', '')}\n"
                        f"Показания одометра: {report_data.data.get('Показания одометра', '')}\n"
                        f"Прибытие на заявку: {report_data.data.get('Прибытие на заявку', '')} комплекс {report_data.data.get('Комплекс', '')}\n"
                        f"Время прибытия: {report_data.data.get('Время прибытия', '')}\n"
                        f"Предпринятые действия: {', '.join(report_data.data.get('Предпринятые действия', ''))}\n"
                        f"Время убытия: {report_data.data.get('Время убытия', '')}\n"
                        f"Время пребывания на месте работы: {report_data.data.get('Время пребывания на месте работы', '')} часа\n"
                        f"Промежуточные показания одометра: {report_data.data.get('Промежуточные показания одометра', '')}\n"
                        f"Показания одометра: {report_data.data.get('Показания одометра', '')}\n"
                        f"Пройдено расстояния всего: {report_data.data.get('Пройдено расстояния всего', '')} км."
                        )


class FillReport(StatesGroup):
    time_departure = State()  # Время выезда
    odometer_reading = State()  # Показания одометра
    take_picture_new_odometer = State() #Фотография начального показателя одометра
    arrival_object = State() #Номер заявки
    arrival_complex = State() # Номер или навание обьекта
    arrival_time = State()  # Время прибытие на заявку
    actions_taken = State()  # Предпринятые действия
    departure_time = State()  # Время убытия
    intermediate_odometer = State()  # Промежуточные показания одометра
    take_picture_intermediate_odometer = State()  # Фотография промежуточного показателя одометра
    final_odometer = State()  # Показания одометра
    take_picture_final_odometer = State()  # Фотография финального показателя одометра
    total_distance = State() # Пройдено км
    time_spent = State() # проведено времени на объекте


# Обработчик для команды /fill_report, начинающей процесс заполнения отчета
@dp.message(Command("fill_report"))
async def fill_report_start(message: Message, state: FSMContext):
    # Очищаем состояние, если оно было сохранено
    await state.clear()
    await state.set_state(FillReport.time_departure)
    await message.answer("Введите время выезда (например, '9 35').")


# Обработчики для заполнения данных отчета
@dp.message(FillReport.time_departure)
async def process_time_departure(message: Message, state: FSMContext):
    await state.update_data(time_departure=message.text)
    await state.set_state(FillReport.odometer_reading)
    await message.answer("Отправьте показания одометра (например, '250465' обязательно 6 знаков).")


@dp.message(FillReport.odometer_reading)
async def process_odometer_reading(message: Message, state: FSMContext):
    await state.update_data(odometer_reading=message.text)
    await state.set_state(FillReport.take_picture_new_odometer)
    await message.answer(
        "Отправьте фотографию с начальным показанием одометра")


@dp.message(FillReport.take_picture_new_odometer, F.photo)
async def process_picture_odometer1(message: Message, state: FSMContext):
    await state.update_data(take_picture_new_odometer=message.photo[-1].file_id)
    await state.set_state(FillReport.arrival_object)
    await message.answer(
        "Отправьте номер заявки (например, '176085').")


@dp.message(FillReport.arrival_object)
async def process_arrival_object(message: Message, state: FSMContext):
    await state.update_data(arrival_object=message.text)
    await state.set_state(FillReport.arrival_complex)
    await message.answer(
        "Отправьте название объекта или "
        "комплекса (например,'LBS11400-LBL11401').")


@dp.message(FillReport.arrival_complex)
async def process_arrival_complex(message: Message, state: FSMContext):
    await state.update_data(arrival_complex=message.text)
    await state.set_state(FillReport.arrival_time)
    await message.answer(
        "Отправьте время прибытия на объект (например, '10 15').")


@dp.message(FillReport.arrival_time)
async def process_arrival_time(message: Message, state: FSMContext):
    await state.update_data(arrival_time=message.text)
    await state.set_state(FillReport.actions_taken)
    await message.answer(
        "Отправьте предпринятые действия (например, 'перенастройка, перезапуск комплекса').")


@dp.message(FillReport.actions_taken)
async def process_actions_taken(message: Message, state: FSMContext):
    await state.update_data(actions_taken=message.text)
    await state.set_state(FillReport.departure_time)
    await message.answer(
        "Отправьте время убытия (например, '12 30').")


@dp.message(FillReport.departure_time)
async def process_departure_time(message: Message, state: FSMContext):
    await state.update_data(departure_time=message.text)
    await state.set_state(FillReport.intermediate_odometer)
    await message.answer(
        "Отправьте промежуточные показания одометра (например, '250556' обязательно 6 знаков).")


@dp.message(FillReport.intermediate_odometer)
async def process_intermediate_odometer(message: Message, state: FSMContext):
    await state.update_data(intermediate_odometer=message.text)
    await state.set_state(FillReport.take_picture_intermediate_odometer)
    await message.answer(
        "Отправьте фотографию с промежуточными показаниями одометра.")


@dp.message(FillReport.take_picture_intermediate_odometer, F.photo)
async def process_picture_odometer2(message: Message, state: FSMContext):
    await state.update_data(take_picture_intermediate_odometer=message.photo[-1].file_id)
    await state.set_state(FillReport.final_odometer)
    await message.answer(
        "Отправьте конечные показания одометра (например, '250556' обязательно 6 знаков).")


@dp.message(FillReport.final_odometer)
async def process_final_odometer(message: Message, state: FSMContext):
    await state.update_data(final_odometer=message.text)
    await state.set_state(FillReport.take_picture_final_odometer)
    await message.answer(
        "Отправьте фотографию с конечными показаниями одометра.")


@dp.message(FillReport.take_picture_final_odometer, F.photo)
async def process_picture_odometer3(message: Message, state: FSMContext):
    # Получаем данные из контекста состояния
    state_data = await state.get_data()
    print("State data:", state_data)  # Выводим содержимое state_data
    # Обновляем данные состояния
    await state.update_data(take_picture_final_odometer=message.photo[-1].file_id)
    await state.update_data(total_distance=total_distance(
        state_data.get('odometer_reading'),
        state_data.get('final_odometer')
    ))
    await state.update_data(time_spent=time_odds(
        state_data.get('arrival_time'),
        state_data.get('departure_time')
    ))
    # Получаем обновленные данные
    data = await state.get_data()
    await state.clear()
    # Создаем объект отчета и формируем текст сообщения
    report_data = ReportData()
    report_data.fill_report(data)
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent} часа\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км."
    )
    # Отправляем сообщение с данными
    await message.answer(report_text)



# Запуск бота
async def main():
    await bot.delete_webhook(drop_pending_updates=True)
    await dp.start_polling(bot)


if __name__ == '__main__':
    asyncio.run(main())



