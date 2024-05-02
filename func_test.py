import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import locale


class ReportData:
    def __init__(self, text):
        self.data = {}
        self.extract_data(text)

    def extract_data(self, text):
        data = {}
        data['Время выезда'] = re.findall(r'время выезда (\d+ \d+)', text)[0]
        data['Показания одометра'] = re.findall(r'показания одометра (\d+)', text)[0]
        data['Прибытие на заявку'] = re.findall(r'прибытие на заявку (\d+)', text)[0]
        data['Комплекс'] = re.findall(r'комплекс (\w+-\w+)', text)[0]
        data['Время прибытия'] = re.findall(r'время прибытия (\d+ \d+)', text)[0]
        predp_deistviya = re.findall(r'предпринятые действия (.+?) время убытия', text)
        data['Предпринятые действия'] = predp_deistviya[0].split() if predp_deistviya else []
        data['Время убытия'] = re.findall(r'время убытия (\d+ \d+)', text)[0]
        data['Время пребывания на месте работы'] = re.findall(r'время пребывания на месте работы (\d+)', text)[0]
        data['Промежуточные показания одометра'] = re.findall(r'промежуточные показания одометра (\d+)', text)[0]
        data['Показания одометра'] = re.findall(r'показания одометра (\d+)', text)[1]
        data['Пройдено расстояния всего'] = re.findall(r'пройдено расстояния всего (\d+)', text)[0]

        self.data = data


def create_report(report_data, report_date):
    # Создание нового документа
    document = Document()

    # Настройки документа
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal'].font.size = Pt(14)

    # Заголовок
    title = document.add_paragraph()
    title_run1 = title.add_run('Отчет по АВР')
    title_run1.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run1.style.font.size = Pt(16)
    title.paragraph_format.space_after = Pt(0)  # Установка интервала "После" в 0pt

    title2 = document.add_paragraph()
    title_run2 = title2.add_run(f'за {report_date}.')
    title_run2.bold = True
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run2.style.font.size = Pt(16)

    # Добавление пустой строки
    document.add_paragraph()

    # Добавление данных из класса ReportData
    start_time = document.add_paragraph()
    start_time.add_run('Время выезда: ').bold = True
    start_time.add_run(report_data.data.get('Время выезда', ''))
    start_time.paragraph_format.space_after = Pt(0)

    start_odometer = document.add_paragraph()
    start_odometer.add_run('Показания одометра: ').bold = True
    start_odometer.add_run(report_data.data.get('Показания одометра', ''))
    start_odometer.paragraph_format.space_after = Pt(0)

    pribytie_text = document.add_paragraph()
    pribytie_text.add_run('Прибытие на заявку: ').bold = True
    pribytie_text.add_run(f"{report_data.data.get('Прибытие на заявку', '')} комплекс {report_data.data.get('Комплекс', '')}")
    pribytie_text.paragraph_format.space_after = Pt(0)

    object_time = document.add_paragraph()
    object_time.add_run('Время прибытия: ').bold = True
    object_time.add_run(report_data.data.get('Время прибытия', ''))
    object_time.paragraph_format.space_after = Pt(0)

    action = document.add_paragraph()
    action.add_run('Предпринятые действия: ').bold = True
    action.add_run(', '.join(report_data.data.get('Предпринятые действия', '')))
    action.paragraph_format.space_after = Pt(0)

    final_time = document.add_paragraph()
    final_time.add_run('Время убытия: ').bold = True
    final_time.add_run(report_data.data.get('Время убытия', ''))
    final_time.paragraph_format.space_after = Pt(0)

    cost_time = document.add_paragraph()
    cost_time.add_run('Время пребывания на месте работы: ').bold = True
    cost_time.add_run(f"{report_data.data.get('Время пребывания на месте работы', '')} часа")
    cost_time.paragraph_format.space_after = Pt(0)

    inter_odometer = document.add_paragraph()
    inter_odometer.add_run('Промежуточные показания одометра: ').bold = True
    inter_odometer.add_run(report_data.data.get('Промежуточные показания одометра', ''))
    inter_odometer.paragraph_format.space_after = Pt(0)

    final_odometer = document.add_paragraph()
    final_odometer.add_run('Показания одометра: ').bold = True
    final_odometer.add_run(report_data.data.get('Показания одометра', ''))
    final_odometer.paragraph_format.space_after = Pt(0)

    distance = document.add_paragraph()
    distance.add_run('Пройдено расстояния всего: ').bold = True
    distance.add_run(f"{report_data.data.get('Пройдено расстояния всего', '')} км.")
    distance.paragraph_format.space_after = Pt(0)

    # Сохранение документа
    document.save(f'Отчет по АВР на {report_date}.docx')


text = "время выезда 9 35 показания одометра 250465 прибытие на заявку 176085 комплекс LBS11400-LBL11401 время прибытия 11 30 Предпринятые действия перезагрузка комплекса перенаведение комплекса время убытия 13 30 время пребывания на месте работы 2 часа промежуточные показания одометра 250556 прибытие на базу показания одометра 250647 пройдено расстояния всего 182 километра"

report_data = ReportData(text)
report_date = datetime.datetime.now().strftime("%d %m %Y")  # текущая дата в формате "дд месяц год"

create_report(report_data, report_date)

