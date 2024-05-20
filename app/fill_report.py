import os
from aiogram import types, F, Router
from aiogram.types import Message, FSInputFile
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from aiogram.filters.command import Command
from app.keyboards import main_kb
import re
import datetime

routers = Router()


def find_first_match(pattern, text, default="Нет данных"):
    matches = re.findall(pattern, text)
    return matches[0] if matches else default


class ReportData:
    def __init__(self):
        self.time_departure = None
        self.odometer_reading = None
        self.take_picture_new_odometer = None
        self.arrival_object = None
        self.arrival_complex = None
        self.arrival_time = None
        self.actions_taken = None
        self.departure_time = None
        self.intermediate_odometer = None
        self.take_picture_intermediate_odometer = None
        self.final_odometer = None
        self.take_picture_final_odometer = None
        self.total_distance = None
        self.time_spent = None

    def count_filled_pictures(self):
        """
        Подсчитывает количество заполненных картинок.
        """
        filled = 0
        if self.take_picture_new_odometer:
            filled += 1
        if self.take_picture_intermediate_odometer:
            filled += 1
        if self.take_picture_final_odometer:
            filled += 1
        return filled

    def count_remaining_pictures(self):
        """
        Подсчитывает количество оставшихся картинок для заполнения.
        Возвращает кортеж (количество оставшихся картинок, список имен картинок для заполнения).
        """
        remaining = 3 - self.count_filled_pictures()
        remaining_pictures = []
        if not self.take_picture_new_odometer:
            remaining_pictures.append("Новый одометр")
        if not self.take_picture_intermediate_odometer:
            remaining_pictures.append("Промежуточный одометр")
        if not self.take_picture_final_odometer:
            remaining_pictures.append("Конечный одометр")
        return remaining, remaining_pictures

    def fill_report(self, data):
        # Обновляем данные только если они присутствуют в переданном словаре
        for key, value in data.items():
            setattr(self, key, value)

    def get_report(self):
        # Возвращает отчет в виде словаря
        return {
            'time_departure': self.time_departure,
            'odometer_reading': self.odometer_reading,
            'take_picture_new_odometer': self.take_picture_new_odometer,
            'arrival_object': self.arrival_object,
            'arrival_complex': self.arrival_complex,
            'arrival_time': self.arrival_time,
            'actions_taken': self.actions_taken,
            'departure_time': self.departure_time,
            'intermediate_odometer': self.intermediate_odometer,
            'take_picture_intermediate_odometer': self.take_picture_intermediate_odometer,
            'final_odometer': self.final_odometer,
            'take_picture_final_odometer': self.take_picture_final_odometer,
            # Добавьте остальные переменные, если есть
        }


user_report_data = {}


def normalize_time(time_str):
    # Паттерн для времени в различных форматах
    time_pattern = re.compile(r'(\d+)\D+(\d+)')

    # Пытаемся найти соответствие времени в строке
    match = time_pattern.search(time_str)
    if match:
        # Если соответствие найдено, извлекаем часы и минуты
        hours, minutes = match.groups()
        return f"{hours.zfill(2)}:{minutes.zfill(2)}"  # Форматируем часы и минуты к двузначному виду (с нулём впереди)
    else:
        return None  # Возвращаем None, если соответствие не найдено


def normalize_odometer(reading_str):
    # Убираем все символы, кроме цифр
    reading_digits = re.sub(r'\D', '', reading_str)
    return reading_digits.zfill(6)  # Добавляем нули впереди, чтобы получить шестизначное число


async def extract_data(user_id, text):
    global user_report_data
    report_data = user_report_data.setdefault(user_id, ReportData())
    data = {}

    time_departure_match = re.findall(r'время выезда (.*?) показания одометра', text)
    if time_departure_match:
        normalized_time = normalize_time(time_departure_match[0])
        data['time_departure'] = normalized_time if normalized_time else None
    else:
        data['time_departure'] = None

    odometer_reading_match = re.findall(r'начальные показания одометра (.*?) прибытие на заявку', text)
    if odometer_reading_match:
        normalized_reading = normalize_odometer(odometer_reading_match[0])
        data['odometer_reading'] = normalized_reading if normalized_reading else None
    else:
        data['odometer_reading'] = None

    arrival_object_match = re.findall(r'прибытие на заявку (.*?) комплекс', text)
    if arrival_object_match:
        data['arrival_object'] = arrival_object_match[0].strip()
    else:
        data['arrival_object'] = None

    arrival_complex_match = re.findall(r'комплекс (.*?) время прибытия на заявку', text)
    if arrival_complex_match:
        data['arrival_complex'] = arrival_complex_match[0].strip()
    else:
        data['arrival_complex'] = None

    arrival_time_match = re.findall(r'время прибытия на заявку (.*?) предпринятые действия', text)
    if arrival_time_match:
        normalized_time = normalize_time(arrival_time_match[0])
        data['arrival_time'] = normalized_time if normalized_time else None
    else:
        data['arrival_time'] = None

    actions_taken_match = re.findall(r'предпринятые действия (.*?) время убытия', text, re.DOTALL)
    if actions_taken_match:
        data['actions_taken'] = actions_taken_match[0].strip()
    else:
        data['actions_taken'] = None

    departure_time_match = re.findall(r'время убытия (.*?) промежуточные показания одометра', text)
    if departure_time_match:
        normalized_time = normalize_time(departure_time_match[0])
        data['departure_time'] = normalized_time if normalized_time else None
    else:
        data['departure_time'] = None

    intermediate_odometer_match = (re.findall
                                   (r'промежуточные показания одометра (.*?) конечные показания одометра', text))
    if intermediate_odometer_match:
        normalized_reading = normalize_odometer(intermediate_odometer_match[0])
        data['intermediate_odometer'] = normalized_reading if normalized_reading else None
    else:
        data['intermediate_odometer'] = None

    final_odometer_match = re.findall(r'конечные показания одометра (.*?) ', text)
    if final_odometer_match:
        normalized_reading = normalize_odometer(final_odometer_match[0])
        data['final_odometer'] = normalized_reading if normalized_reading else None
    else:
        data['final_odometer'] = None

    # Присваивание значений в объект класса ReportData
    report_data.time_departure = data['time_departure']
    report_data.odometer_reading = data['odometer_reading']
    report_data.arrival_object = data['arrival_object']
    report_data.arrival_complex = data['arrival_complex']
    report_data.arrival_time = data['arrival_time']
    report_data.actions_taken = data['actions_taken']
    report_data.departure_time = data['departure_time']
    report_data.time_spent = time_odds(
        report_data.arrival_time,
        report_data.departure_time
    )
    report_data.intermediate_odometer = data['intermediate_odometer']
    report_data.final_odometer = data['final_odometer']
    report_data.total_distance = total_distance(
        report_data.odometer_reading,
        report_data.final_odometer
    )

    # Формирование текста отчета
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Начальные показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent}\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Конечные показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
        f"\n"
        f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
        f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
        f"\n"
        f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
        f"\n"
        f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
    )

    # Отправка сообщения с отчетом
    return report_text


def total_distance(start_reading, end_reading):
    try:
        # Используем регулярное выражение для извлечения только цифр из строки
        start_reading = re.sub(r'\D', '', start_reading)
        end_reading = re.sub(r'\D', '', end_reading)

        # Проверка, что входные значения можно преобразовать в числа
        start_reading = int(start_reading)
        end_reading = int(end_reading)

        # Расчет общего пройденного расстояния
        distance = end_reading - start_reading

        return distance
    except (TypeError, ValueError):
        return None


def time_odds(start_time, end_time):
    def plural_form(number, one, two, five):
        if number % 10 == 1 and number % 100 != 11:
            return one
        elif 2 <= number % 10 <= 4 and (number % 100 < 10 or number % 100 >= 20):
            return two
        else:
            return five

    try:
        # Приведение формата времени к общему виду (часы:минуты)
        start_time = re.sub(r'[^\d:]', '', start_time)
        end_time = re.sub(r'[^\d:]', '', end_time)

        start_hour, start_minute = map(int, start_time.split(':'))
        end_hour, end_minute = map(int, end_time.split(':'))
        total_hours = end_hour - start_hour
        total_minutes = end_minute - start_minute
        if total_minutes < 0:
            total_hours -= 1
            total_minutes += 60

        hour_str = plural_form(total_hours, "час", "часа", "часов")
        minute_str = plural_form(total_minutes, "минута", "минуты", "минут")

        return f"{total_hours} {hour_str} {total_minutes} {minute_str}"
    except (AttributeError, ValueError, TypeError):
        return None


# Обработчик текстового сообщения с данными для отчета
@routers.message(Command("view_report"))
async def process_report_text(message: types.Message):
    global user_report_data
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Начальные показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent}\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Конечные показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
        f"\n"
        f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
        f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
        f"\n"
        f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
        f"\n"
        f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
    )
    await message.answer(report_text, reply_markup=main_kb)


# Объявление состояний
class FillReport(StatesGroup):
    time_departure = State()  # Время выезда
    odometer_reading = State()  # Показания одометра
    take_picture_new_odometer = State()  # Фотография начального показателя одометра
    arrival_object = State()  # Номер заявки
    arrival_complex = State()  # Номер или название объекта
    arrival_time = State()  # Время прибытия на заявку
    actions_taken = State()  # Предпринятые действия
    departure_time = State()  # Время убытия
    intermediate_odometer = State()  # Промежуточные показания одометра
    take_picture_intermediate_odometer = State()  # Фотография промежуточного показателя одометра
    final_odometer = State()  # Показания одометра
    take_picture_final_odometer = State()  # Фотография финального показателя одометра
    total_distance = State()  # Пройдено км
    time_spent = State()  # проведено времени на объекте


# Обработчик для команды /fill_report, начинающей процесс заполнения отчета
@routers.message(Command("fill_report"))
async def fill_report_start(message: Message, state: FSMContext):
    await state.set_state(FillReport.time_departure)
    await message.answer("Введите время выезда (например, '9:35').\nИли напишите /skip для пропуска шага, /cancel для отмены.")


# Обработчики для заполнения данных отчета
@routers.message(FillReport.time_departure)
async def process_time_departure(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.odometer_reading)
        await message.answer("Этап пропущен. Отправьте показания одометра (например, '250465' обязательно 6 знаков).")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    elif re.match(r"\d{1,2}:\d{2}", message.text):
        await state.update_data(time_departure=message.text)
        await state.set_state(FillReport.odometer_reading)
        await message.answer("Отправьте показания одометра (например, '250465' обязательно 6 знаков).")
    else:
        await message.answer("Введите корректное значение времени в формате 'чч:мм'.")


@routers.message(FillReport.odometer_reading)
async def process_odometer_reading(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.take_picture_new_odometer)
        await message.answer("Этап пропущен. Отправьте фотографию с начальным показанием одометра.")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    elif message.text.isdigit() and len(message.text) == 6:
        await state.update_data(odometer_reading=message.text)
        await state.set_state(FillReport.take_picture_new_odometer)
        await message.answer("Отправьте фотографию с начальным показанием одометра.")
    else:
        await message.answer("Введите корректное значение одометра, состоящее из 6 цифр.")


@routers.message(FillReport.take_picture_new_odometer, F.photo)
async def process_picture_odometer1(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.arrival_object)
        await message.answer("Этап пропущен. Отправьте номер заявки (например, '176085').")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    else:
        await state.update_data(take_picture_new_odometer=message.photo[-1].file_id)
        await state.set_state(FillReport.arrival_object)
        await message.answer("Отправьте номер заявки (например, '176085').")


@routers.message(FillReport.arrival_object)
async def process_arrival_object(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.arrival_complex)
        await message.answer("Этап пропущен. Отправьте название объекта или комплекса (например, 'LBS11400-LBL11401').")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    else:
        await state.update_data(arrival_object=message.text)
        await state.set_state(FillReport.arrival_complex)
        await message.answer("Отправьте название объекта или комплекса (например, 'LBS11400-LBL11401').")


@routers.message(FillReport.arrival_complex)
async def process_arrival_complex(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.arrival_time)
        await message.answer("Этап пропущен. Отправьте время прибытия на объект (например, '10:15').")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    else:
        await state.update_data(arrival_complex=message.text)
        await state.set_state(FillReport.arrival_time)
        await message.answer("Отправьте время прибытия на объект (например, '10:15').")


@routers.message(FillReport.arrival_time)
async def process_arrival_time(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.actions_taken)
        await message.answer("Этап пропущен. Отправьте предпринятые действия (например, 'перенастройка, перезапуск комплекса').")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    elif re.match(r"\d{1,2}:\d{2}", message.text):
        await state.update_data(arrival_time=message.text)
        await state.set_state(FillReport.actions_taken)
        await message.answer("Отправьте предпринятые действия (например, 'перенастройка, перезапуск комплекса').")
    else:
        await message.answer("Введите корректное значение времени в формате 'чч:мм'.")


@routers.message(FillReport.actions_taken)
async def process_actions_taken(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.departure_time)
        await message.answer("Этап пропущен. Отправьте время убытия (например, '12:30').")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    else:
        await state.update_data(actions_taken=message.text)
        await state.set_state(FillReport.departure_time)
        await message.answer("Отправьте время убытия (например, '12:30').")


@routers.message(FillReport.departure_time)
async def process_departure_time(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.intermediate_odometer)
        await message.answer("Этап пропущен. Отправьте промежуточные показания одометра (например, '250556' обязательно 6 знаков).")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    elif re.match(r"\d{1,2}:\d{2}", message.text):
        await state.update_data(departure_time=message.text)
        await state.set_state(FillReport.intermediate_odometer)
        await message.answer("Отправьте промежуточные показания одометра (например, '250556' обязательно 6 знаков).")
    else:
        await message.answer("Введите корректное значение времени в формате 'чч:мм'.")


@routers.message(FillReport.intermediate_odometer)
async def process_intermediate_odometer(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.take_picture_intermediate_odometer)
        await message.answer("Этап пропущен. Отправьте фотографию с промежуточными показаниями одометра.")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    elif message.text.isdigit() and len(message.text) == 6:
        await state.update_data(intermediate_odometer=message.text)
        await state.set_state(FillReport.take_picture_intermediate_odometer)
        await message.answer("Отправьте фотографию с промежуточными показаниями одометра.")
    else:
        await message.answer("Введите корректное значение одометра, состоящее из 6 цифр.")


@routers.message(FillReport.take_picture_intermediate_odometer, F.photo)
async def process_picture_odometer2(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.final_odometer)
        await message.answer("Этап пропущен. Отправьте конечные показания одометра (например, '250556' обязательно 6 знаков).")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    else:
        await state.update_data(take_picture_intermediate_odometer=message.photo[-1].file_id)
        await state.set_state(FillReport.final_odometer)
        await message.answer("Отправьте конечные показания одометра (например, '250556' обязательно 6 знаков).")


@routers.message(FillReport.final_odometer)
async def process_final_odometer(message: Message, state: FSMContext):
    if message.text == "/skip":
        await state.set_state(FillReport.take_picture_final_odometer)
        await message.answer("Этап пропущен. Отправьте фотографию с конечными показаниями одометра.")
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    elif message.text.isdigit() and len(message.text) == 6:
        await state.update_data(final_odometer=message.text)
        await state.set_state(FillReport.take_picture_final_odometer)
        await message.answer("Отправьте фотографию с конечными показаниями одометра.")
    else:
        await message.answer("Введите корректное значение одометра, состоящее из 6 цифр.")


@routers.message(FillReport.take_picture_final_odometer, F.photo)
async def process_picture_odometer3(message: Message, state: FSMContext):
    if message.text == "/skip":
        await finalize_report(message, state)
    elif message.text == "/cancel":
        await state.clear()
        await message.answer("Процесс заполнения отчета отменен.")
    else:
        global user_report_data  # Объявляем глобальную переменную
        user_id = message.from_user.id  # Получаем ID пользователя
        state_data = await state.get_data()  # Получаем данные из состояния FSM

        # Обновляем состояние FSM с идентификатором файла фотографии
        await state.update_data(take_picture_final_odometer=message.photo[-1].file_id)

        # Получаем или создаем объект ReportData для пользователя
        report_data = user_report_data.setdefault(user_id, ReportData())
        await message.bot.download(file=message.photo[-1].file_id)
        # Обновляем объект ReportData полученными данными
        report_data.take_picture_final_odometer = message.photo[-1].file_id
        report_data.time_spent = time_odds(
            state_data.get('arrival_time'),
            state_data.get('departure_time')
        )
        report_data.total_distance = total_distance(
            state_data.get('odometer_reading'),
            state_data.get('final_odometer')
        )
        report_data.fill_report(state_data)
        report_text = (
            f"Отчет по АВР:\n"
            f"Время выезда: {report_data.time_departure}\n"
            f"Начальные показания одометра: {report_data.odometer_reading}\n"
            f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
            f"Время прибытия: {report_data.arrival_time}\n"
            f"Предпринятые действия: {report_data.actions_taken}\n"
            f"Время убытия: {report_data.departure_time}\n"
            f"Время пребывания на месте работы: {report_data.time_spent}\n"
            f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
            f"Конечные показания одометра: {report_data.final_odometer}\n"
            f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
            f"\n"
            f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
            f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
            f"\n"
        )
        await message.answer(report_text, reply_markup=main_kb)
        await state.clear()


async def finalize_report(message: Message, state: FSMContext, user_id):
    state_data = await state.get_data()
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {state_data.get('time_departure', 'не указано')}\n"
        f"Начальные показания одометра: {state_data.get('odometer_reading', 'не указано')}\n"
        f"Прибытие на заявку: {state_data.get('arrival_object', 'не указано')} комплекс {state_data.get('arrival_complex', 'не указано')}\n"
        f"Время прибытия: {state_data.get('arrival_time', 'не указано')}\n"
        f"Предпринятые действия: {state_data.get('actions_taken', 'не указано')}\n"
        f"Время убытия: {state_data.get('departure_time', 'не указано')}\n"
        f"Промежуточные показания одометра: {state_data.get('intermediate_odometer', 'не указано')}\n"
        f"Конечные показания одометра: {state_data.get('final_odometer', 'не указано')}\n"
        f"Фотографии:\n"
        f"Начальные показания: {state_data.get('take_picture_new_odometer', 'не указано')}\n"
        f"Промежуточные показания: {state_data.get('take_picture_intermediate_odometer', 'не указано')}\n"
        f"Конечные показания: {state_data.get('take_picture_final_odometer', 'не указано')}\n"
    )
    await message.answer(report_text)
    await state.clear()


@routers.message(FillReport.take_picture_final_odometer, F.photo)
async def process_picture_odometer3(message: Message, state: FSMContext):
    global user_report_data  # Объявляем глобальную переменную
    user_id = message.from_user.id  # Получаем ID пользователя
    state_data = await state.get_data()  # Получаем данные из состояния FSM

    # Обновляем состояние FSM с идентификатором файла фотографии
    await state.update_data(take_picture_final_odometer=message.photo[-1].file_id)

    # Получаем или создаем объект ReportData для пользователя
    report_data = user_report_data.setdefault(user_id, ReportData())
    await message.bot.download(file=message.photo[-1].file_id)
    # Обновляем объект ReportData полученными данными
    report_data.take_picture_final_odometer = message.photo[-1].file_id
    report_data.time_spent = time_odds(
        state_data.get('arrival_time'),
        state_data.get('departure_time')
    )
    report_data.total_distance = total_distance(
        state_data.get('odometer_reading'),
        state_data.get('final_odometer')
    )
    report_data.fill_report(state_data)
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Начальные показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent}\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Конечные показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
        f"\n"
        f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
        f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
        f"\n"
    )
    await message.answer(report_text, reply_markup=main_kb)
    await state.clear()


@routers.message(F.text.lower() == "создать отчет")
async def create_report(message: Message):
    # Создание нового документа
    from main import bot
    report_date = datetime.datetime.now().strftime("%d %m %Y")
    user_id = message.from_user.id
    document = Document()
    global user_report_data
    report_data = user_report_data.setdefault(user_id, ReportData())
    if report_data.count_filled_pictures() < 3:
        await message.reply("Невозможно создать отчет. Необходимо заполнить три картинки.")
        report_text = (
            f"Отчет по АВР:\n"
            f"Время выезда: {report_data.time_departure}\n"
            f"Начальные показания одометра: {report_data.odometer_reading}\n"
            f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
            f"Время прибытия: {report_data.arrival_time}\n"
            f"Предпринятые действия: {report_data.actions_taken}\n"
            f"Время убытия: {report_data.departure_time}\n"
            f"Время пребывания на месте работы: {report_data.time_spent}\n"
            f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
            f"Конечные показания одометра: {report_data.final_odometer}\n"
            f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
            f"\n"
            f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
            f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
            f"\n"
            f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
            f"\n"
            f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
        )
        await message.answer(report_text, reply_markup=main_kb)
    else:
        # Настройки документа
        document.styles['Normal'].font.name = 'Times New Roman'
        document.styles['Normal'].font.size = Pt(14)

        # Заголовок
        title = document.add_paragraph()
        title_run1 = title.add_run('Отчет по АВР')
        title_run1.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run1.style.font.size = Pt(16)
        title.paragraph_format.space_after = Pt(0)  # Установка интервала "После" в 0pt

        title2 = document.add_paragraph()
        title_run2 = title2.add_run(f'за {report_date}.')
        title_run2.bold = True
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run2.style.font.size = Pt(16)
        # Добавление пустой строки
        document.add_paragraph()
        # Добавление данных из класса ReportData
        start_time = document.add_paragraph()
        start_time.add_run('Время выезда: ').bold = True
        start_time.add_run(report_data.time_departure)
        start_time.paragraph_format.space_after = Pt(0)

        start_odometer = document.add_paragraph()
        start_odometer.add_run('Показания одометра: ').bold = True
        start_odometer.add_run(report_data.odometer_reading)
        start_odometer.paragraph_format.space_after = Pt(0)
        # скачивание картинки с начальным одометром
        file_pic_new_odometer_info = await bot.get_file(report_data.take_picture_new_odometer)
        file_pic_new_odometer_download = await bot.download_file(file_pic_new_odometer_info.file_path)
        temp_new_odometer = f"./picture/{report_data.take_picture_new_odometer}.png"

        with open(temp_new_odometer, 'wb') as f:
            f.write(file_pic_new_odometer_download.read())

        picture_new_odometer = document.add_paragraph()
        document.add_picture(temp_new_odometer, width=Inches(3))
        picture_new_odometer.paragraph_format.space_after = Pt(0)

        pribytie_text = document.add_paragraph()
        pribytie_text.add_run('Прибытие на заявку: ').bold = True
        pribytie_text.add_run(
            f"{report_data.arrival_object} комплекс {report_data.arrival_complex}")
        pribytie_text.paragraph_format.space_after = Pt(0)

        object_time = document.add_paragraph()
        object_time.add_run('Время прибытия: ').bold = True
        object_time.add_run(report_data.arrival_time)
        object_time.paragraph_format.space_after = Pt(0)

        action = document.add_paragraph()
        action.add_run('Предпринятые действия: ').bold = True
        action.add_run(report_data.actions_taken)
        action.paragraph_format.space_after = Pt(0)

        final_time = document.add_paragraph()
        final_time.add_run('Время убытия: ').bold = True
        final_time.add_run(report_data.departure_time)
        final_time.paragraph_format.space_after = Pt(0)

        cost_time = document.add_paragraph()
        cost_time.add_run('Время пребывания на месте работы: ').bold = True
        cost_time.add_run(report_data.time_spent)
        cost_time.paragraph_format.space_after = Pt(0)

        inter_odometer = document.add_paragraph()
        inter_odometer.add_run('Промежуточные показания одометра: ').bold = True
        inter_odometer.add_run(report_data.intermediate_odometer)
        inter_odometer.paragraph_format.space_after = Pt(0)

        # скачивание картинки с промежуточным одометром
        file_pic_int_odometer_info = await bot.get_file(report_data.take_picture_intermediate_odometer)
        file_pic_int_odometer_download = await bot.download_file(file_pic_int_odometer_info.file_path)
        temp_int_odometer = f"./picture/{report_data.take_picture_intermediate_odometer}.png"

        with open(temp_int_odometer, 'wb') as f:
            f.write(file_pic_int_odometer_download.read())

        picture_int_odometer = document.add_paragraph()
        document.add_picture(temp_int_odometer, width=Inches(3))
        picture_int_odometer.paragraph_format.space_after = Pt(0)

        final_odometer = document.add_paragraph()
        final_odometer.add_run('Конечные показания одометра: ').bold = True
        final_odometer.add_run(report_data.final_odometer)
        final_odometer.paragraph_format.space_after = Pt(0)

        # скачивание картинки с конечным одометром
        file_pic_fin_odometer_info = await bot.get_file(report_data.take_picture_final_odometer)
        file_pic_fin_odometer_download = await bot.download_file(file_pic_fin_odometer_info.file_path)
        temp_fin_odometer = f"./picture/{report_data.take_picture_final_odometer}.png"

        with open(temp_fin_odometer, 'wb') as f:
            f.write(file_pic_fin_odometer_download.read())

        picture_fin_odometer = document.add_paragraph()
        document.add_picture(temp_fin_odometer, width=Inches(3))
        picture_fin_odometer.paragraph_format.space_after = Pt(0)

        distance = document.add_paragraph()
        distance.add_run('Пройдено расстояния всего: ').bold = True
        distance.add_run(str(report_data.total_distance))
        distance.paragraph_format.space_after = Pt(0)
        # Сохранение документа
        document.save(f'./docs/Отчет по АВР на {report_date}.docx')
        document = FSInputFile(path=f'./docs/Отчет по АВР на {report_date}.docx')
        await message.answer_document(document=document)
        os.remove(temp_new_odometer)
        os.remove(temp_int_odometer)
        os.remove(temp_fin_odometer)
        os.remove(f'./docs/Отчет по АВР на {report_date}.docx')


class EditTimeDeparture(StatesGroup):
    time_departure = State()


@routers.message(F.text.lower() == "изменить время выезда")
async def start_edit_time_departure(message: Message, state: FSMContext):
    await state.set_state(EditTimeDeparture.time_departure)
    await message.answer("Введите новое время выезда")


@routers.message(EditTimeDeparture.time_departure)
async def finish_edit_time_departure(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    # Получаем данные пользователя из словаря. Если они уже есть, используем их, иначе создаем новые.
    report_data = user_report_data.setdefault(user_id, ReportData())
    if re.match(r"\d{1,2}:\d{2}", message.text):
        report_data.time_departure = message.text
        report_text = (
            f"Отчет по АВР:\n"
            f"Время выезда: {report_data.time_departure}\n"
            f"Начальные показания одометра: {report_data.odometer_reading}\n"
            f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
            f"Время прибытия: {report_data.arrival_time}\n"
            f"Предпринятые действия: {report_data.actions_taken}\n"
            f"Время убытия: {report_data.departure_time}\n"
            f"Время пребывания на месте работы: {report_data.time_spent}\n"
            f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
            f"Конечные показания одометра: {report_data.final_odometer}\n"
            f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
            f"\n"
            f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
            f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
            f"\n"
            f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
            f"\n"
            f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
        )
        # Отправляем обновленный отчет
        await message.answer(report_text, reply_markup=main_kb)
        # Очищаем состояние
        await state.clear()
    else:
        await message.answer("Введите корректное значение времени в формате 'чч:мм'.")


class EditOdometerReading(StatesGroup):
    odometer_reading = State()


@routers.message(F.text.lower() == "изменить значение начального одометра")
async def start_edit_odometer_reading(message: Message, state: FSMContext):
    await state.set_state(EditOdometerReading.odometer_reading)
    await message.answer("Введите новое значение начального одометра")


@routers.message(EditOdometerReading.odometer_reading)
async def finish_edit_odometer_reading(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    # Получаем данные пользователя из словаря. Если они уже есть, используем их, иначе создаем новые.
    report_data = user_report_data.setdefault(user_id, ReportData())
    if message.text.isdigit() and len(message.text) == 6:
        report_data.odometer_reading = message.text  # Изменяем время выезда
        report_data.total_distance = total_distance(
            report_data.odometer_reading,
            report_data.final_odometer
        )
        report_text = (
            f"Отчет по АВР:\n"
            f"Время выезда: {report_data.time_departure}\n"
            f"Начальные показания одометра: {report_data.odometer_reading}\n"
            f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
            f"Время прибытия: {report_data.arrival_time}\n"
            f"Предпринятые действия: {report_data.actions_taken}\n"
            f"Время убытия: {report_data.departure_time}\n"
            f"Время пребывания на месте работы: {report_data.time_spent}\n"
            f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
            f"Конечные показания одометра: {report_data.final_odometer}\n"
            f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
            f"\n"
            f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
            f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
            f"\n"
            f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
            f"\n"
            f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
        )
        await message.answer(report_text, reply_markup=main_kb)
        await state.clear()
    else:
        await message.answer("Введите корректное значение одометра, состоящее из 6 цифр.")


class EditOdometerReadingPicture(StatesGroup):
    take_picture_new_odometer = State()


@routers.message(F.text.lower() == "изменить фото начального одометра")
async def start_edit_picture_odometer(message: Message, state: FSMContext):
    await state.set_state(EditOdometerReadingPicture.take_picture_new_odometer)
    await message.answer("Отправьте новое фото с начальным одометром")


@routers.message(EditOdometerReadingPicture.take_picture_new_odometer, F.photo)
async def finish_edit_picture_odometer(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    report_data.take_picture_new_odometer = message.photo[-1].file_id
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Начальные показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent}\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Конечные показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
        f"\n"
        f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
        f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
        f"\n"
        f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
        f"\n"
        f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
    )
    await message.answer(report_text, reply_markup=main_kb)
    await state.clear()


class EditArrivalObject(StatesGroup):
    arrival_object = State()


@routers.message(F.text.lower() == "изменить номер заявки")
async def start_edit_arrival_object(message: Message, state: FSMContext):
    await state.set_state(EditArrivalObject.arrival_object)
    await message.answer("Введите новый номер заявки")


@routers.message(EditArrivalObject.arrival_object)
async def finish_edit_arrival_object(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    report_data.arrival_object = message.text
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Начальные показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent}\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Конечные показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
        f"\n"
        f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
        f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
        f"\n"
        f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
        f"\n"
        f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
    )
    await message.answer(report_text, reply_markup=main_kb)
    await state.clear()


class EditArrivalComplex(StatesGroup):
    arrival_complex = State()


@routers.message(F.text.lower() == "изменить название комплекса")
async def start_edit_arrival_complex(message: Message, state: FSMContext):
    await state.set_state(EditArrivalComplex.arrival_complex)
    await message.answer("Введите новое название комплекса")


@routers.message(EditArrivalComplex.arrival_complex)
async def finish_edit_arrival_complex(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    report_data.arrival_object = message.text
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Начальные показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent}\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Конечные показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
        f"\n"
        f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
        f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
        f"\n"
        f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
        f"\n"
        f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
    )
    await message.answer(report_text, reply_markup=main_kb)
    await state.clear()


class EditArrivalTime(StatesGroup):
    arrival_time = State()


@routers.message(F.text.lower() == "изменить время прибытия на объект")
async def start_edit_arrival_time(message: Message, state: FSMContext):
    await state.set_state(EditArrivalTime.arrival_time)
    await message.answer("Введите новое время прибытия на объект")


@routers.message(EditArrivalTime.arrival_time)
async def finish_edit_arrival_time(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    if re.match(r"\d{1,2}:\d{2}", message.text):
        report_data.arrival_time = message.text
        report_data.time_spent = time_odds(
            report_data.arrival_time,
            report_data.departure_time
        )
        report_text = (
            f"Отчет по АВР:\n"
            f"Время выезда: {report_data.time_departure}\n"
            f"Начальные показания одометра: {report_data.odometer_reading}\n"
            f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
            f"Время прибытия: {report_data.arrival_time}\n"
            f"Предпринятые действия: {report_data.actions_taken}\n"
            f"Время убытия: {report_data.departure_time}\n"
            f"Время пребывания на месте работы: {report_data.time_spent}\n"
            f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
            f"Конечные показания одометра: {report_data.final_odometer}\n"
            f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
            f"\n"
            f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
            f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
            f"\n"
            f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
            f"\n"
            f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
        )
        await message.answer(report_text, reply_markup=main_kb)
        await state.clear()
    else:
        await message.answer("Введите корректное значение времени в формате 'чч:мм'.")


class EditActionsTaken(StatesGroup):
    actions_taken = State()


@routers.message(F.text.lower() == "изменить действия")
async def start_edit_actions_taken(message: Message, state: FSMContext):
    await state.set_state(EditActionsTaken.actions_taken)
    await message.answer("Введите предпринятые действия")


@routers.message(EditActionsTaken.actions_taken)
async def finish_edit_actions_taken(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    report_data.actions_taken = message.text
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Начальные показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent}\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Конечные показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
        f"\n"
        f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
        f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
        f"\n"
        f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
        f"\n"
        f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
    )
    await message.answer(report_text, reply_markup=main_kb)
    await state.clear()


class EditDepartureTime(StatesGroup):
    departure_time = State()


@routers.message(F.text.lower() == "изменить время убытия")
async def start_edit_departure_time(message: Message, state: FSMContext):
    await state.set_state(EditDepartureTime.departure_time)
    await message.answer("Введите новое время убытия с объекта")


@routers.message(EditDepartureTime.departure_time)
async def finish_edit_departure_time(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    if re.match(r"\d{1,2}:\d{2}", message.text):
        report_data.departure_time = message.text
        report_data.time_spent = time_odds(
            report_data.arrival_time,
            report_data.departure_time
        )
        report_text = (
            f"Отчет по АВР:\n"
            f"Время выезда: {report_data.time_departure}\n"
            f"Начальные показания одометра: {report_data.odometer_reading}\n"
            f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
            f"Время прибытия: {report_data.arrival_time}\n"
            f"Предпринятые действия: {report_data.actions_taken}\n"
            f"Время убытия: {report_data.departure_time}\n"
            f"Время пребывания на месте работы: {report_data.time_spent}\n"
            f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
            f"Конечные показания одометра: {report_data.final_odometer}\n"
            f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
            f"\n"
            f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
            f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
            f"\n"
            f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
            f"\n"
            f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
        )
        await message.answer(report_text, reply_markup=main_kb)
        await state.clear()
    else:
        await message.answer("Введите корректное значение времени в формате 'чч:мм'.")


class EditIntermediateOdometer(StatesGroup):
    intermediate_odometer = State()


@routers.message(F.text.lower() == "изменить значения промеж. одометра")
async def start_edit_intermediate_odometer(message: Message, state: FSMContext):
    await state.set_state(EditIntermediateOdometer.intermediate_odometer)
    await message.answer("Введите новое значение промежуточного одометра")


@routers.message(EditIntermediateOdometer.intermediate_odometer)
async def finish_edit_intermediate_odometer(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    if message.text.isdigit() and len(message.text) == 6:
        report_data.intermediate_odometer = message.text  # Изменяем время выезда
        report_text = (
            f"Отчет по АВР:\n"
            f"Время выезда: {report_data.time_departure}\n"
            f"Начальные показания одометра: {report_data.odometer_reading}\n"
            f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
            f"Время прибытия: {report_data.arrival_time}\n"
            f"Предпринятые действия: {report_data.actions_taken}\n"
            f"Время убытия: {report_data.departure_time}\n"
            f"Время пребывания на месте работы: {report_data.time_spent}\n"
            f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
            f"Конечные показания одометра: {report_data.final_odometer}\n"
            f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
            f"\n"
            f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
            f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
            f"\n"
            f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
            f"\n"
            f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
        )
        await message.answer(report_text, reply_markup=main_kb)
        await state.clear()
    else:
        await message.answer("Введите корректное значение одометра, состоящее из 6 цифр.")


class EditPictureIntermediateOdometer(StatesGroup):
    take_picture_intermediate_odometer = State()


@routers.message(F.text.lower() == "изменить фото промеж. одометра")
async def start_edit_picture_intermediate_odometer(message: Message, state: FSMContext):
    await state.set_state(EditPictureIntermediateOdometer.take_picture_intermediate_odometer)
    await message.answer("Отправьте новое фото с промежуточным одометром")


@routers.message(EditPictureIntermediateOdometer.take_picture_intermediate_odometer, F.photo)
async def finish_edit_picture_intermediate_odometer(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    report_data.take_picture_intermediate_odometer = message.photo[-1].file_id
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Начальные показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent}\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Конечные показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
        f"\n"
        f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
        f"Осталось заполнить картинок: {report_data.count_remaining_pictures()}\n"
        f"\n"
        f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
    )
    await message.answer(report_text, reply_markup=main_kb)
    await state.clear()


class EditFinalOdometer(StatesGroup):
    final_odometer = State()


@routers.message(F.text.lower() == "изменить конечное значение одометра")
async def start_edit_final_odometer(message: Message, state: FSMContext):
    await state.set_state(EditFinalOdometer.final_odometer)
    await message.answer("Введите новое конечное значение одометра")


@routers.message(EditFinalOdometer.final_odometer)
async def finish_edit_final_odometer(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    if len(message.text) == 6 and message.text.isdigit():
        report_data.final_odometer = message.text  # Изменяем конечное значение одометра
        report_data.total_distance = total_distance(
            report_data.odometer_reading,
            report_data.final_odometer
        )
        report_text = (
            f"Отчет по АВР:\n"
            f"Время выезда: {report_data.time_departure}\n"
            f"Начальные показания одометра: {report_data.odometer_reading}\n"
            f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
            f"Время прибытия: {report_data.arrival_time}\n"
            f"Предпринятые действия: {report_data.actions_taken}\n"
            f"Время убытия: {report_data.departure_time}\n"
            f"Время пребывания на месте работы: {report_data.time_spent}\n"
            f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
            f"Конечные показания одометра: {report_data.final_odometer}\n"
            f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
            f"\n"
            f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
            f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
            f"\n"
            f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
            f"\n"
            f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
        )
        await message.answer(report_text, reply_markup=main_kb)
        await state.clear()
    else:
        await message.answer("Введите корректное значение одометра, состоящее из 6 цифр.")


class EditPictureFinalOdometer(StatesGroup):
    take_picture_final_odometer = State()


@routers.message(F.text.lower() == "изменить фото конечного одометра")
async def start_edit_picture_final_odometer(message: Message, state: FSMContext):
    await state.set_state(EditPictureFinalOdometer.take_picture_final_odometer)
    await message.answer("Отправьте новое фото с конечным одометром")


@routers.message(EditPictureFinalOdometer.take_picture_final_odometer, F.photo)
async def finish_edit_picture_intermediate_odometer(message: Message, state: FSMContext):
    global user_report_data  # Указываем, что используем глобальную переменную
    user_id = message.from_user.id
    report_data = user_report_data.setdefault(user_id, ReportData())
    report_data.take_picture_final_odometer = message.photo[-1].file_id
    report_text = (
        f"Отчет по АВР:\n"
        f"Время выезда: {report_data.time_departure}\n"
        f"Начальные показания одометра: {report_data.odometer_reading}\n"
        f"Прибытие на заявку: {report_data.arrival_object} комплекс {report_data.arrival_complex}\n"
        f"Время прибытия: {report_data.arrival_time}\n"
        f"Предпринятые действия: {report_data.actions_taken}\n"
        f"Время убытия: {report_data.departure_time}\n"
        f"Время пребывания на месте работы: {report_data.time_spent}\n"
        f"Промежуточные показания одометра: {report_data.intermediate_odometer}\n"
        f"Конечные показания одометра: {report_data.final_odometer}\n"
        f"Пройдено расстояния всего: {report_data.total_distance} км.\n"
        f"\n"
        f"Заполнено картинок: {report_data.count_filled_pictures()}\n"
        f"Осталось заполнить картинок: {report_data.count_remaining_pictures()[0]}\n"
        f"\n"
        f"Необходимо заполнить следующие картинки: {', '.join(report_data.count_remaining_pictures()[1])}\n"
        f"\n"
        f"Не забудьте дополнить или изменить отчет, добавить картинки через клавиатуру!"
    )
    await message.answer(report_text, reply_markup=main_kb)
    await state.clear()
